# File: reporter.py

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
from visualizer import plot_bar, plot_box, plot_heatmap, plot_volcano
from utils import sanitize_filename

def create_report(analyzer, stats_list, p_values_list, posthoc_list, results_text_list, ci_list, plot_dir, report_path, corr_matrix=None, corr_narrative=None, volcano_results=None, volcano_narratives=None, control_group=None):
    doc = Document()
    doc.add_heading('Advanced Experimental Results Report', level=0)
    doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    for i, stats_df in enumerate(stats_list):
        parameter = stats_df['Parameter'].iloc[0]
        safe_param_name = sanitize_filename(parameter)
        doc.add_heading(parameter, level=1) 
        
        table = doc.add_table(rows=1, cols=len(stats_df.columns))
        table.style = 'Table Grid' 
        hdr_cells = table.rows[0].cells
        for j, col in enumerate(stats_df.columns):
            hdr_cells[j].text = str(col)
        for _, row in stats_df.iterrows():
            row_cells = table.add_row().cells
            for j, val in enumerate(row):
                row_cells[j].text = str(val)

        bar_path = os.path.join(plot_dir, f'{safe_param_name}_bar.png')
        plot_bar(analyzer.data, parameter, analyzer.group_col, bar_path, p_values_list[i], posthoc_list[i], ci_list[i])
        doc.add_picture(bar_path, width=Inches(5))

        box_path = os.path.join(plot_dir, f'{safe_param_name}_box.png')
        plot_box(analyzer.data, parameter, analyzer.group_col, box_path, p_values_list[i], posthoc_list[i])
        doc.add_picture(box_path, width=Inches(5))

        doc.add_paragraph(results_text_list[i])

    if corr_matrix is not None:
        doc.add_heading("Parameter Correlation", level=1)
        heatmap_path = os.path.join(plot_dir, "correlation_heatmap.png")
        plot_heatmap(corr_matrix, heatmap_path)
        doc.add_picture(heatmap_path, width=Inches(6))
        if corr_narrative: doc.add_paragraph(corr_narrative)

    # --- NEW: Volcano Plot Section ---
    if volcano_results and control_group:
        doc.add_heading("Global Treatment Effects (Volcano Plots)", level=1)
        doc.add_paragraph(f"Visualizing global parameter shifts relative to the '{control_group}' baseline.")
        
        for trt, v_df in volcano_results.items():
            doc.add_heading(f"{trt} vs {control_group}", level=2)
            
            # Save and insert plot
            volcano_path = os.path.join(plot_dir, f"volcano_{sanitize_filename(trt)}.png")
            plot_volcano(v_df, trt, control_group, volcano_path)
            doc.add_picture(volcano_path, width=Inches(6))
            
            # Insert automated narrative
            if trt in volcano_narratives:
                doc.add_paragraph(volcano_narratives[trt])

    try:
        doc.save(report_path)
        print(f"✅ Success! Advanced report saved to: {report_path}")
    except PermissionError:
        print(f"\n❌ ERROR: Permission Denied. Please close the document in Word and run again!")